from docx import Document
import json

def extract_document_structure(doc_path):
    """Extract the structure and content of a Word document"""
    doc = Document(doc_path)
    
    # Extract document properties
    properties = {
        'title': doc.core_properties.title,
        'author': doc.core_properties.author,
        'created': str(doc.core_properties.created) if doc.core_properties.created else None,
    }
    
    # Extract paragraphs with their styles
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            paragraphs.append({
                'text': para.text,
                'style': para.style.name,
                'alignment': para.alignment,
                'bold': any(run.bold for run in para.runs),
                'italic': any(run.italic for run in para.runs),
            })
    
    # Extract tables
    tables = []
    for table in doc.tables:
        t = []
        for row in table.rows:
            r = []
            for cell in row.cells:
                r.append(cell.text)
            t.append(r)
        tables.append(t)
    
    # Extract sections
    sections = []
    for section in doc.sections:
        sections.append({
            'orientation': section.orientation,
            'page_width': section.page_width.inches,
            'page_height': section.page_height.inches,
            'margins': {
                'top': section.top_margin.inches,
                'bottom': section.bottom_margin.inches,
                'left': section.left_margin.inches,
                'right': section.right_margin.inches,
            }
        })
    
    return {
        'properties': properties,
        'paragraphs': paragraphs,
        'tables': tables,
        'sections': sections,
    }

if __name__ == "__main__":
    doc_path = "1- JOCC-2204-1016 ammar[1] (AutoRecovered).docx"
    structure = extract_document_structure(doc_path)
    
    # Print document structure in a readable format
    print("Document Properties:")
    for key, value in structure['properties'].items():
        print(f"  {key}: {value}")
    
    print("\nParagraphs:")
    for i, para in enumerate(structure['paragraphs']):
        print(f"  {i+1}. Style: {para['style']}, Bold: {para['bold']}, Italic: {para['italic']}")
        print(f"     Text: {para['text'][:100]}..." if len(para['text']) > 100 else f"     Text: {para['text']}")
    
    print(f"\nTables: {len(structure['tables'])}")
    
    print("\nSections:")
    for i, section in enumerate(structure['sections']):
        print(f"  Section {i+1}: {section['orientation']}, "
              f"{section['page_width']}x{section['page_height']} inches")
    
    # Save full structure to a JSON file for further analysis
    with open('document_structure.json', 'w') as f:
        json.dump(structure, f, indent=2)
    
    print("\nFull document structure saved to document_structure.json")